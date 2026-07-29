"""
Narrative and supply-chain sections for the Word report.

Split out of docx_report.py so the sections that consume the structured
narrative (see narrative.py) and the deterministic supply-chain view (see
supply_chain.py) can grow without turning build_docx into one long function.

Every function here takes a Kit: the document plus the formatting helpers
build_docx already defines, so the two files cannot drift apart on fonts,
colours or table styling.

Nothing in this module computes a figure. Counts come from supply_chain.py,
which groups what Koi reported; prose comes from the narrative, whose
citations were verified before it got here.
"""

from __future__ import annotations

from dataclasses import dataclass
from typing import Any, Callable


@dataclass
class Kit:
    """The document and the formatting helpers, handed over by build_docx."""

    doc: Any
    heading: Callable
    lede: Callable
    note: Callable
    table: Callable
    para: Callable
    palette: dict


def _sev_colour(kit: Kit, level: str) -> str | None:
    return kit.palette.get(str(level).lower())


def _pct(value: int, total: int) -> str:
    return f"{100 * value / total:.1f}%" if total else "\u2014"


# ------------------------------------------------------------ supply chain


def supply_chain(kit: Kit, view: dict) -> None:
    """Where software comes from, and what is known about the parties behind it."""
    if not view:
        return
    channels = view.get("channels") or []
    dimensions = (view.get("findings") or {}).get("dimensions") or []
    if not channels and not dimensions:
        return

    kit.doc.add_page_break()
    kit.heading("Software supply chain")
    kit.lede("Where this software comes from, and what is known about the "
             "parties behind it.")

    profile = view.get("publishers") or {}
    ratio = profile.get("items_per_publisher")
    if profile.get("publishers"):
        kit.para(
            f"The estate carries {profile['items']:,} items sourced from "
            f"{profile['publishers']:,} distinct third-party publishers, an "
            f"average of {ratio:.1f} items per publisher. Each publisher is a "
            "separate trust decision, and each is a party whose compromise "
            "would reach this estate through software you already run."
        )

    if channels:
        kit.heading("How software enters the estate", 2)
        kit.lede("Marketplaces and registries grouped by the route they "
                 "represent. Each item is counted once, at its source.")
        total = sum(c["total"] for c in channels)
        kit.table(
            ["Channel", "Sources", "Items", "Share"],
            [[c["label"],
              ", ".join(f"{n} ({q:,})" for n, q in c["members"])[:90],
              format(c["total"], ","),
              _pct(c["total"], total)]
             for c in channels],
        )
        kit.note(
            "Shares are of the " + format(total, ",") + " items whose source "
            "the platform recorded, so they total 100%. Items with no recorded "
            "marketplace are absent from this table and from those shares."
        )

    if dimensions:
        kit.heading("Trust signals by dimension", 2)
        kit.lede("Findings reported by Koi, grouped by the supply-chain "
                 "question each one answers. One item can carry several "
                 "findings, so these are counts of findings, not of items.")
        for dim in dimensions:
            _dimension(kit, dim)

    uncategorised = (view.get("findings") or {}).get("uncategorised") or []
    if uncategorised:
        kit.heading("Not yet categorised", 2)
        kit.lede("Finding types with no supply-chain dimension yet. Listed "
                 "rather than dropped, so nothing is understated.")
        kit.table(["Finding", "Occurrences"],
                  [[str(r["finding"]).replace("_", " "), format(r["count"], ",")]
                   for r in uncategorised[:10]])


def _dimension(kit: Kit, dim: dict) -> None:
    kit.heading(dim["label"], 2)
    mitre = ", ".join(dim.get("mitre") or [])
    kit.lede(dim["question"] + (f"  \u00b7  MITRE {mitre}" if mitre else ""))

    seen = dim.get("items_seen") or 0
    if dim.get("understated"):
        kit.para(
            f"{seen} of the ranked items carry a finding in this dimension. "
            "This snapshot did not record an estate-wide count for it, so no "
            "total is shown rather than a misleading zero."
        )
    elif dim.get("total"):
        kit.para(
            f"{dim['total']:,} findings across the estate; {seen} of the "
            "ranked items carry one."
        )
    else:
        kit.para("No finding in this dimension was reported for this estate.")
        return

    if dim.get("findings"):
        kit.table(["Finding", "Occurrences"],
                  [[str(r["finding"]).replace("_", " "), format(r["count"], ",")]
                   for r in dim["findings"][:8]])

    if dim.get("examples"):
        kit.table(
            ["Item", "Publisher", "Source", "Finding", "Endpoints"],
            [[(e.get("name") or "")[:34],
              (e.get("publisher") or "")[:24],
              (e.get("marketplace") or "")[:18],
              ", ".join(e.get("matched") or [])[:44],
              format(int(e.get("endpoints") or 0), ",")]
             for e in dim["examples"]],
        )


# ----------------------------------------------------------------- findings


def findings(kit: Kit, blocks: list[dict]) -> None:
    """The analyst's read of the estate, each claim carrying its citations."""
    if not blocks:
        return
    kit.doc.add_page_break()
    kit.heading("Findings")
    kit.lede("Each finding cites the collected data it rests on. Confidence "
             "is stated explicitly.")

    for block in blocks:
        kit.heading(block["title"], 2)
        _tags(kit, [block["severity"].upper(), block["confidence"].upper()]
              + list(block.get("mitre_techniques") or []),
              _sev_colour(kit, block["severity"]))
        if block.get("narrative"):
            kit.para(block["narrative"])
        if block.get("affected_scope"):
            kit.para("Scope: " + block["affected_scope"])
        _evidence(kit, block.get("evidence"))


def scenarios(kit: Kit, blocks: list[dict]) -> None:
    """Paths an attacker could take. Illustrative, and labelled as such."""
    if not blocks:
        return
    kit.doc.add_page_break()
    kit.heading("Attack scenarios")
    kit.lede("Paths an attacker could take given the exposure observed in "
             "this tenant. These are illustrative, not incidents that have "
             "occurred.")

    for block in blocks:
        kit.heading(block["title"], 2)
        _tags(kit, [block["likelihood"].upper()]
              + list(block.get("mitre_techniques") or []), None)
        for step in block["steps"]:
            kit.doc.add_paragraph(step, style="List Number")
        if block.get("impact"):
            kit.para("Impact. " + block["impact"])
        if block.get("breaks_at"):
            kit.para("What breaks this chain. " + block["breaks_at"])
        _evidence(kit, block.get("enabling_evidence"))


def actions(kit: Kit, blocks: list[dict], fallback: str = "") -> None:
    """Recommended actions, ordered by risk reduced rather than by ease."""
    kit.doc.add_page_break()
    kit.heading("Recommended actions")
    kit.lede("Ordered by risk reduced, not by ease of implementation.")

    if not blocks:
        if fallback:
            for line in fallback.split("\n"):
                if line.strip():
                    kit.doc.add_paragraph(line.strip(), style="List Bullet")
        else:
            kit.para("[[TO BE PROVIDED: recommended actions]]")
        return

    for position, block in enumerate(blocks, start=1):
        kit.heading(f"{position}. {block['title']}", 2)
        _tags(kit, [block["effort"].upper() + " EFFORT"], None)
        if block.get("rationale"):
            kit.para(block["rationale"])
        if block.get("expected_outcome"):
            kit.para(block["expected_outcome"])
        if block.get("platform_capability"):
            kit.note("Platform capability: " + block["platform_capability"])


def threat_context(kit: Kit, blocks: list[dict]) -> None:
    """Public threat activity, fenced off from anything tenant-derived."""
    if not blocks:
        return
    kit.doc.add_page_break()
    kit.heading("Threat context")
    kit.lede("Related public threat activity.")
    kit.note(
        "This section was NOT verified against your tenant. It draws on the "
        "analyst's knowledge of publicly reported activity and is provided for "
        "context only. Every preceding section is traced to data collected "
        "from your environment."
    )
    for block in blocks:
        kit.heading(block["campaign_or_pattern"], 2)
        kit.para(block["relevance"])
        _evidence(kit, block.get("tenant_link"), label="WHAT LINKS IT HERE")


def data_gaps(kit: Kit, gaps: list[str]) -> None:
    """Where the data was thin. Stated plainly rather than left implicit."""
    if not gaps:
        return
    kit.heading("Data gaps", 2)
    kit.lede("What this report could not establish, and why.")
    for gap in gaps:
        kit.doc.add_paragraph(gap, style="List Bullet")


def validation_note(kit: Kit, validation: dict | None) -> None:
    """State how much of the narrative survived citation checking.

    Rendered in the report on purpose: a customer document that claims its
    figures are traceable should say how that was verified.
    """
    if not validation or not validation.get("checked_citations"):
        return
    checked = validation["checked_citations"]
    verified = validation.get("verified_citations", 0)
    dropped = len(validation.get("dropped") or [])
    line = (
        f"Citation check: {verified} of {checked} citations in the analytical "
        "sections were traced back to the collected data."
    )
    if dropped:
        line += (
            f" {dropped} claim(s) could not be traced and were removed rather "
            "than shipped."
        )
    kit.note(line)


# ------------------------------------------------------------------ helpers


def _tags(kit: Kit, labels: list[str], accent: str | None) -> None:
    from docx.shared import Pt, RGBColor

    labels = [label for label in labels if label]
    if not labels:
        return
    paragraph = kit.doc.add_paragraph()
    for position, label in enumerate(labels):
        run = paragraph.add_run(("  \u00b7  " if position else "") + label)
        run.bold = True
        run.font.size = Pt(8)
        hexcode = accent if position == 0 and accent else kit.palette.get("muted")
        if hexcode:
            run.font.color.rgb = RGBColor.from_string(hexcode)


def _evidence(kit: Kit, entries: list[dict] | None, label: str = "EVIDENCE") -> None:
    from docx.shared import Pt, RGBColor

    if not entries:
        return
    header = kit.doc.add_paragraph()
    run = header.add_run(label)
    run.bold = True
    run.font.size = Pt(7.5)
    muted = kit.palette.get("muted")
    if muted:
        run.font.color.rgb = RGBColor.from_string(muted)

    for entry in entries:
        kind = str(entry.get("kind", "")).replace("_", " ")
        text = f"{entry.get('reference', '')} ({kind})"
        if entry.get("note"):
            text += " \u2014 " + entry["note"]
        paragraph = kit.doc.add_paragraph(text, style="List Bullet")
        for run in paragraph.runs:
            run.font.size = Pt(8.5)
