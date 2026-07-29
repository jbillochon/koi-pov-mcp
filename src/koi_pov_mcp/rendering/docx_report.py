"""
Editable Word report.

Ported from jbillochon/povplatform (rendering/report.py build_docx): same
print palette, same section order, built natively with python-docx rather
than converted from HTML, because the point of shipping a DOCX is that
someone can edit it.

The analytical sections (supply chain, findings, attack scenarios,
recommended actions, threat context) live in docx_sections.py; this file owns
the document, the palette and the formatting helpers, and hands them over.
"""

from __future__ import annotations

import logging
from pathlib import Path

from . import docx_sections
from .common import NOT_MEASURED, Data

log = logging.getLogger(__name__)

# Print palette
TEAL_DARK = "0B6E4F"
TEAL = "12805C"
AMBER = "B45309"
RED = "B91C1C"
GREY = "5A6B62"
LIGHT = "7A8B82"

RISK_COLOUR = {
    "critical": RED,
    "high": AMBER,
    "medium": "0D9488",
    "low": TEAL,
    "pending": GREY,
}

SECTION_PALETTE = {**RISK_COLOUR, "info": GREY, "muted": LIGHT}

VERDICT_COLOUR = {
    "met": TEAL_DARK,
    "passed": TEAL_DARK,
    "partially met": AMBER,
    "partial": AMBER,
    "not met": RED,
    "failed": RED,
    "not tested": GREY,
    "not measured": GREY,
}


def build_docx(data: dict, out_path: str, narrative: dict | None = None) -> str:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    r = Data(data)
    n = narrative or {}
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    def colour(run, hexcode: str):
        run.font.color.rgb = RGBColor.from_string(hexcode)

    def heading(text: str, level: int = 1):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        if level == 1:
            run.font.size = Pt(18)
            colour(run, TEAL_DARK)
            p.space_before = Pt(18)
            p.space_after = Pt(4)
        else:
            run.font.size = Pt(13)
            colour(run, TEAL)
            p.space_before = Pt(14)
            p.space_after = Pt(3)
        return p

    def lede(text: str):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.italic = True
        colour(run, GREY)
        return p

    def note(text: str):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(8.5)
        run.italic = True
        colour(run, LIGHT)
        return p

    def para(text: str):
        return doc.add_paragraph(text)

    def table(headers, rows, colour_col=None, palette=None):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            cellp = t.rows[0].cells[i].paragraphs[0]
            run = cellp.add_run(h)
            run.bold = True
            run.font.size = Pt(9)
        for row in rows:
            cells = t.add_row().cells
            for i, value in enumerate(row):
                cellp = cells[i].paragraphs[0]
                run = cellp.add_run(str(value))
                run.font.size = Pt(9)
                if colour_col is not None and i == colour_col:
                    hexcode = (palette or RISK_COLOUR).get(str(value).lower())
                    if hexcode:
                        colour(run, hexcode)
                        run.bold = True
        return t

    kit = docx_sections.Kit(
        doc=doc, heading=heading, lede=lede, note=note, table=table,
        para=para, palette=SECTION_PALETTE,
    )

    # ---------------------------------------------------------------- cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Cortex AES")
    run.bold = True
    run.font.size = Pt(32)
    colour(run, TEAL_DARK)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Proof of Value \u2014 software supply chain assessment")
    run.font.size = Pt(15)
    colour(run, TEAL)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lines = [r.customer]
    if r.meta.get("pov_start") and r.meta.get("pov_end"):
        lines.append(str(r.meta["pov_start"]) + " to " + str(r.meta["pov_end"]))
    if r.meta.get("prepared_by"):
        lines.append("Prepared by " + str(r.meta["prepared_by"]))
    run = info.add_run("\n".join(lines))
    run.font.size = Pt(11)
    colour(run, GREY)

    doc.add_page_break()

    # ---------------------------------------------------- executive summary
    heading("Executive summary")
    lede("What the platform found across the estate during this Proof of Value.")

    if n.get("headline"):
        p = doc.add_paragraph()
        run = p.add_run(n["headline"])
        run.bold = True
        run.font.size = Pt(13)
        colour(run, TEAL_DARK)

    if n.get("executive_summary"):
        for chunk in n["executive_summary"].split("\n"):
            if chunk.strip():
                doc.add_paragraph(chunk.strip())
    else:
        doc.add_paragraph(
            "[[TO BE PROVIDED: executive summary]] \u2014 the figures below were "
            "collected from the tenant, but no narrative was supplied."
        )

    table(
        ["Metric", "Value"],
        [
            ["Endpoints covered", r.num("devices_total")],
            ["Active endpoints", r.num("devices_active")],
            ["Items discovered", r.num("items_total")],
            ["Unique publishers", r.num("unique_publishers")],
            ["High and critical risk", r.high_and_critical],
            ["Ungoverned high risk", r.num("ungoverned_high_risk")],
            ["Policies enabled", r.num("policies_enabled")],
            ["Remediations recorded", r.num("remediations_total")],
            ["Alerts raised", r.num("alerts_total")],
            ["Agent sessions observed", r.num("agent_sessions_total")],
        ],
    )

    stage = r.stage
    if stage:
        doc.add_paragraph(
            "Tenant lifecycle stage: " + stage + ". "
            + ("The tenant is running in observation mode: nothing is currently "
               "blocked, allowlisted or automatically remediated."
               if stage == "discovery"
               else "Policies and remediation are active in this tenant.")
        )
    else:
        doc.add_paragraph(
            "Tenant lifecycle stage could not be determined: the governance "
            "domains were not all collected during this PoV."
        )

    if r.missing:
        note(
            "Not collected during this PoV: " + ", ".join(sorted(r.missing))
            + ". Figures for these areas are absent, not zero, and are shown as "
            + NOT_MEASURED + " throughout."
        )

    heading("How this data was collected", 2)
    doc.add_paragraph(
        "Every figure in this report is read from the customer's Koi tenant "
        "through its API. Endpoint inventory is gathered by a script the "
        "tenant's existing EDR or MDM runs periodically, so no additional agent "
        "was deployed. Risk scores come from the Koi risk engine. Threat "
        "intelligence, where present, comes from public sources named in that "
        "section and is dated."
    )
    doc.add_paragraph(
        "The analytical sections that follow were written from that collected "
        "data. Every claim cites the specific items it rests on, and those "
        "citations were verified against the collected data before this report "
        "was generated; claims that could not be traced were removed."
    )

    # ------------------------------------------------------- success criteria
    doc.add_page_break()
    heading("Success criteria")
    lede("What was agreed at kickoff, and what the collected data shows.")
    criteria = n.get("success_criteria") or []
    if criteria:
        table(
            ["Criterion", "Verdict", "Evidence"],
            [[c.get("criterion", ""), str(c.get("verdict", "")).capitalize(),
              c.get("evidence", "")] for c in criteria],
            colour_col=1, palette=VERDICT_COLOUR,
        )
    else:
        doc.add_paragraph(
            "[[TO BE PROVIDED: success criteria agreed at kickoff]] \u2014 none "
            "were supplied, so no verdict can be recorded here. This is the "
            "section the customer will read first; it should not ship empty."
        )

    # ------------------------------------------------------------- discovery
    doc.add_page_break()
    heading("Discovery")
    lede("What is actually installed across the estate.")

    views = r.get("items_by_view")
    if views:
        heading("By category", 2)
        table(["Category", "Items"],
              [[k, format(v, ",")] for k, v in list(views.items())[:12]])
    markets = r.get("items_by_marketplace")
    if markets:
        heading("By marketplace", 2)
        table(["Marketplace", "Items"],
              [[k, format(v, ",")] for k, v in list(markets.items())[:12]])

    risk = r.get("items_by_risk")
    if risk:
        heading("Risk distribution", 2)
        table(
            ["Level", "Items"],
            [[lvl.capitalize(), format(risk[lvl], ",")]
             for lvl in ("critical", "high", "medium", "low", "pending")
             if lvl in risk],
            colour_col=0,
        )
        if risk.get("pending"):
            note(
                format(risk["pending"], ",") + " items are still awaiting risk "
                "analysis. Their eventual scores are not reflected above."
            )
    elif not r.measured("items_by_risk"):
        doc.add_paragraph("Inventory was " + NOT_MEASURED + " during this PoV.")

    # --------------------------------------------------------- supply chain
    docx_sections.supply_chain(kit, n.get("supply_chain") or {})

    # --------------------------------------------------------- risk inventory
    top = r.rows("top_risk_items")
    if top:
        doc.add_page_break()
        heading("Highest-risk items")
        lede("Scored by the Koi risk engine, ordered by score.")
        rows = []
        for item in top[:25]:
            score = item.get("risk")
            rows.append([
                (item.get("name") or "")[:44],
                (item.get("publisher") or "")[:26],
                item.get("marketplace") or "",
                (item.get("risk_level") or "").capitalize(),
                format(score, ".1f") if isinstance(score, (int, float)) else "\u2014",
                format(int(item.get("endpoints") or 0), ","),
            ])
        table(["Item", "Publisher", "Marketplace", "Risk", "Score", "Endpoints"],
              rows, colour_col=3)

    action = r.rows("action_candidates")
    if action:
        heading("Ungoverned high-risk items", 2)
        lede("High or critical items that no policy currently governs, "
             "ordered by installed footprint.")
        table(
            ["Item", "Marketplace", "Risk", "Endpoints", "Findings"],
            [[(i.get("name") or "")[:40], i.get("marketplace") or "",
              (i.get("risk_level") or "").capitalize(),
              format(int(i.get("endpoints") or 0), ","),
              ", ".join((i.get("findings") or [])[:3])[:60]]
             for i in action[:15]],
            colour_col=2,
        )

    # ------------------------------------------------ findings and scenarios
    docx_sections.findings(kit, n.get("key_findings") or [])
    docx_sections.scenarios(kit, n.get("attack_scenarios") or [])

    # ---------------------------------------------------- threat intelligence
    enr = r.enrichment
    if enr:
        doc.add_page_break()
        heading("Threat intelligence")
        fetched = (enr.get("fetched_at") or "")[:10]
        lede(
            "Threat intel as of " + (fetched or "unknown date")
            + ". Read exploitation first: KEV membership, then EPSS "
            "probability, then CVSS severity."
        )
        ti = r.ti_rows(20)
        if ti:
            heading("CVEs in scope", 2)
            table(["CVE", "Known exploited", "EPSS", "CVSS", "Severity"], ti,
                  colour_col=4)
            if not r.kev_count:
                doc.add_paragraph(
                    "No CVE in scope appears in the CISA Known Exploited "
                    "Vulnerabilities catalogue, so none is currently evidenced "
                    "as exploited in the wild. That is a statement about "
                    "available evidence, not a guarantee of safety."
                )
        osv = r.osv_rows(20)
        if osv:
            heading("Vulnerable packages", 2)
            lede("Matched by exact name and installed version against OSV.dev.")
            table(["Package", "Advisories", "Identifiers"], osv)
        mitre = r.mitre_rows(20)
        if mitre:
            heading("ATT&CK techniques from findings", 2)
            lede("Mapped from Koi finding types through a curated mapping; "
                 "findings with no mapping are omitted rather than guessed.")
            table(["Finding", "Technique", "Name"], mitre)
        unmapped = enr.get("unmapped_findings") or []
        if unmapped:
            note(
                str(len(unmapped)) + " finding types have no ATT&CK mapping yet "
                "and are therefore absent from the table above."
            )

    # ------------------------------------------------------------ governance
    doc.add_page_break()
    heading("Governance")
    lede("Policies, guardrails and lists configured during the PoV.")
    policies = r.rows("policies")
    if policies:
        heading("Marketplace policies", 2)
        table(
            ["Policy", "Action", "State", "Groups"],
            [[(p.get("name") or "")[:44], (p.get("action") or "").capitalize(),
              "Enabled" if p.get("enabled") else "Disabled",
              str(p.get("groups")) if p.get("groups") else "All"]
             for p in policies[:20]],
        )
    elif r.measured("policies"):
        doc.add_paragraph("No marketplace policy was configured during this PoV.")

    runtime = r.rows("runtime_policies")
    if runtime:
        heading("Agent runtime policies", 2)
        table(
            ["Policy", "Mode", "Agents", "Rule types"],
            [[(p.get("name") or "")[:40], (p.get("mode") or "").capitalize(),
              ", ".join(p.get("agents") or [])[:34] or "\u2014",
              ", ".join(sorted(set(p.get("rule_types") or [])))[:34] or "\u2014"]
             for p in runtime[:15]],
        )

    heading("Lists and approvals", 2)
    approvals = r.mapping("approvals_by_status")
    table(
        ["Control", "Count"],
        [["Items allowlisted", r.num("allowlist_count")],
         ["Items blocklisted", r.num("blocklist_count")],
         ["Approvals approved", format(int(approvals.get("approved", 0)), ",")],
         ["Approvals pending", format(int(approvals.get("pending", 0)), ",")]],
    )

    # ----------------------------------------------------------- remediation
    doc.add_page_break()
    heading("Remediation")
    lede("Risk actually removed from endpoints during the PoV.")
    by_status = r.get("remediations_by_status")
    if by_status:
        table(["Status", "Count"],
              [[k.capitalize(), format(v, ",")] for k, v in by_status.items()])
    elif not r.measured("remediations_total"):
        doc.add_paragraph("Remediation data was " + NOT_MEASURED + ".")
    else:
        doc.add_paragraph("No remediation was recorded during this window.")

    remediated = r.rows("remediated_items")
    if remediated:
        heading("Items removed from endpoints", 2)
        table(
            ["Item", "Endpoint", "Platform", "Risk", "Reason"],
            [[(i.get("name") or "")[:40], (i.get("hostname") or "")[:24],
              i.get("platform") or "", (i.get("risk_level") or "").capitalize(),
              (i.get("reason") or "")[:34]]
             for i in remediated[:20]],
            colour_col=3,
        )

    alerts = r.get("alerts_by_severity")
    if alerts:
        heading("Alerts raised", 2)
        table(["Severity", "Count"],
              [[str(k).capitalize(), format(v, ",")] for k, v in alerts.items()],
              colour_col=0)

    # -------------------------------------------------------------- agentic
    doc.add_page_break()
    heading("Agentic runtime activity")
    lede("Coding-agent sessions observed and governed on managed endpoints.")
    decisions = r.get("agent_decisions")
    if decisions:
        table(["Decision", "Actions"],
              [[k.capitalize(), format(v, ",")] for k, v in decisions.items()])
    agents = r.get("agents_seen")
    if agents:
        heading("Agents in use", 2)
        table(["Agent", "Sessions"],
              [[k, format(v, ",")] for k, v in list(agents.items())[:10]])
    blocked = r.rows("agent_blocked_examples")
    if blocked:
        heading("Blocked actions", 2)
        lede("Sample from the last 24 hours; the API caps the event window.")
        table(
            ["Agent", "Endpoint", "Action", "Target"],
            [[b.get("agent") or "", (b.get("host") or "")[:24],
              (b.get("action") or "")[:20], (b.get("target") or "")[:50]]
             for b in blocked[:15]],
        )
    if not r.measured("agent_sessions_total"):
        doc.add_paragraph("Agent activity was " + NOT_MEASURED + " during this PoV.")

    # ------------------------------------------------------ xsiam correlation
    x = r.xsiam
    if x:
        doc.add_page_break()
        heading("Cortex XSIAM correlation")
        lede("Where third-party exposure meets detections raised on the same estate.")
        cov = x.get("coverage") or {}
        inc = x.get("incidents") or {}
        table(
            ["Metric", "Value"],
            [["Endpoints known to Koi", format(int(cov.get("koi_devices", 0)), ",")],
             ["Endpoints known to XSIAM", format(int(cov.get("xsiam_endpoints", 0)), ",")],
             ["Present in both", format(int(cov.get("on_both", 0)), ",")],
             ["Koi only", format(int(cov.get("koi_only_count", 0)), ",")],
             ["XSIAM only", format(int(cov.get("xsiam_only_count", 0)), ",")],
             ["XSIAM incidents in window", format(int(inc.get("total", 0)), ",")]],
        )
        doc.add_paragraph(
            "Correlation here means co-occurrence: the same host appearing in "
            "both platforms. It does not establish that discovered software "
            "caused a detection, and this report never claims that it did."
        )
        hosts = inc.get("top_koi_hosts") or []
        if hosts:
            heading("Koi-known endpoints carrying incidents", 2)
            table(["Endpoint", "Incidents"],
                  [[h.get("host", ""), format(int(h.get("incidents", 0)), ",")]
                   for h in hosts[:20]])

    # ---------------------------------------- actions and threat context
    docx_sections.actions(kit, n.get("recommended_actions") or [],
                          fallback=n.get("recommendations") or "")
    docx_sections.threat_context(kit, n.get("threat_context") or [])

    # -------------------------------------------------------------- appendix
    doc.add_page_break()
    heading("Appendix")

    freq = r.rows("finding_frequency")
    if freq:
        heading("Most frequent findings", 2)
        table(["Finding", "Occurrences"],
              [[str(f.get("finding")).replace("_", " "),
                format(int(f.get("count") or 0), ",")] for f in freq[:15]])

    docx_sections.data_gaps(kit, n.get("data_gaps") or [])

    if r.missing:
        heading("Domains not collected", 2)
        for domain in sorted(r.missing):
            doc.add_paragraph(domain, style="List Bullet")
        note("Figures depending on these domains are absent, not zero.")

    warnings = r.raw("warnings", []) or []
    if warnings:
        heading("Collection warnings", 2)
        for w in warnings[:15]:
            doc.add_paragraph(str(w), style="List Bullet")

    heading("About this report", 2)
    note(
        "Generated by koi-pov-mcp. Figures are read from the customer's Koi "
        "tenant at the time of collection. Narrative sections are written by "
        "the consultant from that same data; anything left unwritten appears "
        "as a visible placeholder rather than being invented."
    )
    docx_sections.validation_note(kit, n.get("validation"))

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    log.info("DOCX written to %s", out_path)
    return out_path
